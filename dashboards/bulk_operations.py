import csv
import io
from django.utils import timezone
from django.db import transaction
from accounts.models import User
from exams.models import Exam, Question, Choice, SchoolClass, Subject
from .models import (
    BulkImportJob, BulkExportJob, StudentPerformance
)


class BulkImporter:
    """Handle bulk imports from CSV"""
    
    @staticmethod
    def import_students(csv_file, school_class, created_by):
        """Import students from CSV, Excel, or Word"""
        job = BulkImportJob.objects.create(
            import_type='students',
            csv_file=csv_file,
            created_by=created_by,
            status='processing',
            started_at=timezone.now(),
        )

        errors = []
        successes = 0

        try:
            if hasattr(csv_file, 'seek'):
                csv_file.seek(0)

            filename = csv_file.name.lower()
            if filename.endswith('.xlsx'):
                rows, parse_errors = BulkImporter._parse_excel_students(csv_file)
            elif filename.endswith('.docx'):
                rows, parse_errors = BulkImporter._parse_word_students(csv_file)
            else:
                rows, parse_errors = BulkImporter._parse_csv_students(csv_file)
            errors.extend(parse_errors)

            for row_num, row in enumerate(rows, start=2):
                try:
                    username = row.get('username', '').strip()
                    email = row.get('email', '').strip()
                    first_name = row.get('first_name', '').strip()
                    last_name = row.get('last_name', '').strip()

                    if not all([username, email, first_name, last_name]):
                        errors.append(f"Row {row_num}: Missing required field")
                        continue

                    if User.objects.filter(username=username).exists():
                        errors.append(f"Row {row_num}: Username '{username}' already exists")
                        continue

                    if User.objects.filter(email=email).exists():
                        errors.append(f"Row {row_num}: Email '{email}' already exists")
                        continue

                    User.objects.create_user(
                        username=username,
                        email=email,
                        first_name=first_name,
                        last_name=last_name,
                        password=row.get('password', '').strip() or 'changeme123',
                        role=User.Role.STUDENT,
                        student_class=school_class,
                        is_approved=True,
                    )
                    successes += 1

                except Exception as e:
                    errors.append(f"Row {row_num}: {str(e)}")

            job.status = 'completed'
            job.total_rows = successes + len(errors)
            job.successful_rows = successes
            job.failed_rows = len(errors)
            job.error_log = '\n'.join(errors) if errors else ''
            job.completed_at = timezone.now()
            job.save()

        except Exception as e:
            job.status = 'failed'
            job.error_log = str(e)
            job.completed_at = timezone.now()
            job.save()

        return job

    @staticmethod
    def _parse_csv_students(file_obj):
        data = []
        errors = []
        try:
            file_content = file_obj.read().decode('utf-8')
            clean_lines = [
                line for line in file_content.splitlines()
                if line.strip() and not line.strip().startswith('#')
            ]
            reader = csv.DictReader(clean_lines)
            required = ['first_name', 'last_name', 'username', 'email']
            if reader.fieldnames and not all(f in reader.fieldnames for f in required):
                return [], [f"Missing required columns: {', '.join(required)}"]
            for row in reader:
                if any(row.values()):
                    data.append(row)
        except Exception as e:
            errors.append(f"CSV Parse Error: {str(e)}")
        return data, errors

    @staticmethod
    def _parse_excel_students(file_obj):
        import openpyxl
        data = []
        errors = []
        try:
            wb = openpyxl.load_workbook(file_obj, data_only=True)
            sheet = wb.active
            headers = [cell.value for cell in sheet[1]]
            required = ['first_name', 'last_name', 'username', 'email']
            if not all(f in headers for f in required):
                return [], [f"Missing required columns: {', '.join(required)}"]
            header_map = {h: i for i, h in enumerate(headers) if h}
            for row in sheet.iter_rows(min_row=2):
                row_data = {}
                has_content = False
                for field, col_idx in header_map.items():
                    val = row[col_idx].value
                    row_data[field] = str(val).strip() if val is not None else ''
                    if val:
                        has_content = True
                if not has_content:
                    continue
                first = row_data.get('first_name', '')
                if first.startswith('#') or first.startswith('•') or first.upper().startswith('NOTE'):
                    continue
                data.append(row_data)
        except Exception as e:
            errors.append(f"Excel Parse Error: {str(e)}")
        return data, errors

    @staticmethod
    def _parse_word_students(file_obj):
        import docx
        data = []
        errors = []
        try:
            doc = docx.Document(file_obj)
            if not doc.tables:
                return [], ["No tables found in Word document"]
            table = doc.tables[0]
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            required = ['first_name', 'last_name', 'username', 'email']
            if not all(f in headers for f in required):
                return [], [f"Missing required columns: {', '.join(required)}"]
            header_map = {h: i for i, h in enumerate(headers) if h}
            for row in table.rows[1:]:
                row_data = {}
                has_content = False
                for field, col_idx in header_map.items():
                    val = row.cells[col_idx].text.strip() if col_idx < len(row.cells) else ''
                    row_data[field] = val
                    if val:
                        has_content = True
                if has_content:
                    data.append(row_data)
        except Exception as e:
            errors.append(f"Word Parse Error: {str(e)}")
        return data, errors
    
    @staticmethod
    def import_questions(csv_file, created_by, exam):
        """Import questions from CSV/Excel/Word directly into an Exam"""
        from exams.models import Question, Choice
        job = BulkImportJob.objects.create(
            import_type='questions',
            csv_file=csv_file,
            created_by=created_by,
            exam=exam,
            status='processing',
            started_at=timezone.now(),
        )

        errors = []
        successes = 0

        try:
            if hasattr(csv_file, 'seek'):
                csv_file.seek(0)

            filename = csv_file.name.lower()
            if filename.endswith('.xlsx'):
                rows, parse_errors = BulkImporter._parse_excel_questions(csv_file)
            elif filename.endswith('.docx'):
                rows, parse_errors = BulkImporter._parse_word_questions(csv_file)
            else:
                rows, parse_errors = BulkImporter._parse_csv_questions(csv_file)
            errors.extend(parse_errors)

            if not rows and not errors:
                errors.append("No data found in file")

            # Determine next order number
            next_order = exam.questions.count() + 1

            for row_num, row in enumerate(rows, start=1):
                try:
                    text = row.get('text', '').strip()
                    q_type = row.get('type', '').lower().strip()
                    marks = int(row.get('marks', '1') or 1)

                    if not text:
                        errors.append(f"Row {row_num}: Missing question text")
                        continue
                    if q_type not in ['objective', 'subjective']:
                        errors.append(f"Row {row_num}: type must be 'objective' or 'subjective', got '{q_type}'")
                        continue

                    question = Question.objects.create(
                        exam=exam,
                        text=text,
                        type=q_type,
                        marks=marks,
                        order=next_order,
                    )
                    next_order += 1

                    if q_type == 'objective':
                        has_correct = False
                        for i in range(1, 5):
                            choice_text = str(row.get(f'choice_{i}', '') or '').strip()
                            is_correct = str(row.get(f'correct_{i}', '') or '').lower().strip() in ['true', '1', 'yes']
                            if choice_text:
                                Choice.objects.create(
                                    question=question,
                                    text=choice_text,
                                    is_correct=is_correct,
                                )
                                if is_correct:
                                    has_correct = True
                        if not has_correct:
                            errors.append(f"Row {row_num}: Objective question has no correct answer marked")

                    successes += 1

                except Exception as e:
                    errors.append(f"Row {row_num}: {str(e)}")

            job.status = 'completed' if successes > 0 or not errors else 'failed'
            job.total_rows = successes + len(errors)
            job.successful_rows = successes
            job.failed_rows = len(errors)
            job.error_log = '\n'.join(errors) if errors else ''
            job.completed_at = timezone.now()
            job.save()

        except Exception as e:
            job.status = 'failed'
            job.error_log = str(e)
            job.completed_at = timezone.now()
            job.save()

        return job

    @staticmethod
    def _parse_csv_questions(file_obj):
        """Parse CSV content, skipping comment rows starting with #"""
        data = []
        errors = []
        try:
            file_content = file_obj.read().decode('utf-8')
            # Strip comment rows before passing to DictReader
            clean_lines = [
                line for line in file_content.splitlines()
                if line.strip() and not line.strip().startswith('#')
            ]
            reader = csv.DictReader(clean_lines)

            required = ['text', 'type', 'marks']
            if reader.fieldnames and not all(f in reader.fieldnames for f in required):
                return [], [f"CSV Missing required fields: {', '.join(required)}"]

            for row in reader:
                if any(row.values()):  # skip fully empty rows
                    data.append(row)
        except Exception as e:
            errors.append(f"CSV Parse Error: {str(e)}")
        return data, errors

    @staticmethod
    def _parse_excel_questions(file_obj):
        """Parse Excel content using openpyxl"""
        import openpyxl
        data = []
        errors = []
        try:
            wb = openpyxl.load_workbook(file_obj, data_only=True)
            sheet = wb.active
            
            # Read header
            headers = [cell.value for cell in sheet[1]]
            required = ['text', 'type', 'marks']
            if not all(field in headers for field in required):
                 return [], [f"Excel Missing required headers: {', '.join(required)}"]

            # Map headers to indices
            header_map = {h: i for i, h in enumerate(headers) if h}

            for row_idx, row in enumerate(sheet.iter_rows(min_row=2), start=2):
                row_data = {}
                has_content = False
                for field, col_idx in header_map.items():
                    val = row[col_idx].value
                    if val is not None:
                        row_data[field] = str(val).strip()
                        has_content = True
                    else:
                        row_data[field] = ''

                if not has_content:
                    continue
                # Skip note/comment rows
                text_val = row_data.get('text', '')
                if text_val.startswith('#') or text_val.startswith('•') or text_val.upper().startswith('NOTE'):
                    continue
                data.append(row_data)
        except Exception as e:
            errors.append(f"Excel Parse Error: {str(e)}")
        return data, errors

    @staticmethod
    def _parse_word_questions(file_obj):
        """Parse Word content using python-docx (Table based)"""
        import docx
        data = []
        errors = []
        try:
            doc = docx.Document(file_obj)
            if not doc.tables:
                return [], ["No tables found in Word document"]
            
            table = doc.tables[0]
            if not table.rows:
                 return [], ["Empty table in Word document"]

            # Header
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            required = ['text', 'type', 'marks']
            if not all(field in headers for field in required):
                 return [], [f"Word Table Missing required headers: {', '.join(required)}"]
            
            header_map = {h: i for i, h in enumerate(headers) if h}

            for row_idx, row in enumerate(table.rows[1:], start=2):
                row_data = {}
                # Ensure row has enough cells
                has_content = False
                for field, col_idx in header_map.items():
                    if col_idx < len(row.cells):
                        val = row.cells[col_idx].text.strip()
                        row_data[field] = val
                        if val: has_content = True
                    else:
                        row_data[field] = ""
                
                if has_content:
                    data.append(row_data)

        except Exception as e:
            errors.append(f"Word Parse Error: {str(e)}")
        return data, errors


class BulkExporter:
    """Handle bulk exports to CSV/Excel"""
    
    @staticmethod
    def export_exam_results(exam, format_type='csv'):
        """Export exam results to CSV/Excel"""
        from exams.models import ExamAttempt
        
        job = BulkExportJob.objects.create(
            export_type='results',
            exam=exam,
            exported_by=None,  # Set by view
            status='processing',
            file_format=format_type,
        )
        
        try:
            attempts = ExamAttempt.objects.filter(
                exam=exam,
                status='submitted'
            ).select_related('student')
            
            # Create CSV in memory
            output = io.StringIO()
            writer = csv.writer(output)
            
            # Header
            writer.writerow([
                'Student ID', 'Student Name', 'Email',
                'Score', 'Total Marks', 'Percentage',
                'Time Taken (mins)', 'Status', 'Submitted At'
            ])
            
            # Data rows
            for attempt in attempts:
                time_minutes = (attempt.completed_at - attempt.started_at).total_seconds() / 60 if attempt.completed_at else 0
                total_marks = sum([q.marks for q in exam.questions.all()])
                percentage = (attempt.total_score / total_marks * 100) if total_marks > 0 else 0
                
                writer.writerow([
                    attempt.student.id,
                    f"{attempt.student.first_name} {attempt.student.last_name}",
                    attempt.student.email,
                    attempt.total_score,
                    total_marks,
                    f"{percentage:.2f}",
                    f"{time_minutes:.2f}",
                    'Passed' if percentage >= 60 else 'Failed',
                    attempt.completed_at.strftime('%Y-%m-%d %H:%M:%S') if attempt.completed_at else '',
                ])
            
            csv_content = output.getvalue()
            
            # Save file
            from django.core.files.base import ContentFile
            filename = f"results_{exam.id}_{timezone.now().strftime('%Y%m%d_%H%M%S')}.csv"
            job.export_file.save(
                filename,
                ContentFile(csv_content.encode('utf-8')),
                save=True
            )
            
            job.status = 'completed'
            job.completed_at = timezone.now()
            job.save()
            
        except Exception as e:
            job.status = 'failed'
            job.completed_at = timezone.now()
            job.save()
        
        return job
    
    @staticmethod
    def export_student_performance(student, format_type='csv'):
        """Export student performance report"""
        job = BulkExportJob.objects.create(
            export_type='performance',
            exported_by=None,  # Set by view
            status='processing',
            file_format=format_type,
        )
        
        try:
            performances = StudentPerformance.objects.filter(
                student=student
            ).select_related('exam').order_by('-attempted_at')
            
            output = io.StringIO()
            writer = csv.writer(output)
            
            writer.writerow([
                'Exam', 'Attempt', 'Score', 'Total Marks',
                'Percentage', 'Time Taken (mins)', 'Status', 'Date'
            ])
            
            for perf in performances:
                time_minutes = perf.time_taken / 60
                writer.writerow([
                    perf.exam.title,
                    perf.attempt_number,
                    perf.score,
                    perf.total_marks,
                    f"{perf.percentage:.2f}",
                    f"{time_minutes:.2f}",
                    perf.status.capitalize(),
                    perf.attempted_at.strftime('%Y-%m-%d %H:%M:%S'),
                ])
            
            csv_content = output.getvalue()
            
            filename = f"performance_{student.id}_{timezone.now().strftime('%Y%m%d_%H%M%S')}.csv"
            from django.core.files.base import ContentFile
            job.export_file.save(
                filename,
                ContentFile(csv_content.encode('utf-8')),
                save=True
            )
            
            job.status = 'completed'
            job.completed_at = timezone.now()
            job.save()
            
        except Exception as e:
            job.status = 'failed'
            job.completed_at = timezone.now()
            job.save()
        
        return job
